# module/commentor.py
from docx import Document
from difflib import SequenceMatcher
import logging, sys
from pathlib import Path
from typing import List, Dict

# ---------------- Console Logging ----------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - [%(levelname)s] - %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)


def find_best_paragraph_match(doc: Document, section_text: str):
    """
    Find the paragraph index in the DOCX most similar to the given section_text.
    """
    best_idx = None
    best_score = 0
    for idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue
        score = SequenceMatcher(None, para_text.lower(), section_text.lower()).ratio()
        if score > best_score:
            best_score = score
            best_idx = idx
    return best_idx, best_score


def add_comments_to_docx(input_file: str, findings: List[Dict], output_file: str):
    """
    Add compliance comments into DOCX file based on findings.
    findings is a list of dicts with keys: section, ai_analysis.
    """
    if not Path(input_file).exists():
        logger.error(f"Input file not found: {input_file}")
        return False

    logger.info(f"Loading DOCX: {input_file}")
    doc = Document(input_file)

    logger.info(f"Adding {len(findings)} compliance comments...")
    for finding in findings:
        section = finding.get("section", "")
        comment_text = finding.get("ai_analysis", "")

        idx, score = find_best_paragraph_match(doc, section)
        if idx is not None and score > 0.3:  # only if match is reasonably good
            para = doc.paragraphs[idx]
            # python-docx doesn't support true comment objects, so we'll append inline markers
            para.add_run(f"  [COMMENT: {comment_text}]").italic = True
        else:
            logger.warning(f"No good match found for section snippet: {section}")

    logger.info(f"Saving annotated DOCX to: {output_file}")
    doc.save(output_file)
    return True


if __name__ == "__main__":
    # Example usage with dummy findings
    dummy_findings = [
        {
            "section": "The company will issue 1 share at incorporation...",
            "ai_analysis": "⚠ Missing clear statement on paid-up capital as per ADGM rules."
        },
        {
            "section": "Resolution authorizing incorporation...",
            "ai_analysis": "⚠ Should explicitly reference ADGM Template Resolution wording."
        }
    ]
    sample_input = "uploaded_docs/sample.docx"
    sample_output = "uploaded_docs/sample_annotated.docx"

    add_comments_to_docx(sample_input, dummy_findings, sample_output)
